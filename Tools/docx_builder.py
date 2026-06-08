"""Build Word documents from Markdown using template styles."""

from __future__ import annotations

import re
from pathlib import Path
from typing import Any

from docx import Document
from docx.table import Table


STYLE_MAP = {
    "h1": "Heading 1",
    "h2": "Heading 2",
    "h3": "Heading 3",
    "paragraph": "Normal",
    "bullet_list": "List Bullet",
    "ordered_list": "List Number",
}


def _get_or_create_styles(doc: Document) -> dict[str, str]:
    """Return available style names in template."""
    available = {s.name for s in doc.styles}
    result = {}
    for key, name in STYLE_MAP.items():
        result[key] = name if name in available else "Normal"
    if "Table Grid" in available:
        result["table"] = "Table Grid"
    else:
        result["table"] = "Normal Table" if "Normal Table" in available else "Normal"
    return result


def _parse_markdown_blocks(body: str) -> list[dict[str, Any]]:
    """Simple line-based markdown block parser."""
    blocks: list[dict[str, Any]] = []
    lines = body.splitlines()
    i = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # Heading
        heading_match = re.match(r"^(#{1,3})\s+(.+)$", stripped)
        if heading_match:
            level = len(heading_match.group(1))
            blocks.append({"type": f"h{level}", "text": heading_match.group(2)})
            i += 1
            continue

        # Table
        if "|" in stripped and i + 1 < len(lines) and re.match(r"^\|?\s*[-:]+", lines[i + 1]):
            table_lines = [stripped]
            i += 1
            i += 1  # skip separator
            while i < len(lines) and "|" in lines[i]:
                table_lines.append(lines[i].strip())
                i += 1
            rows = []
            for tl in table_lines:
                cells = [c.strip() for c in tl.strip("|").split("|")]
                rows.append(cells)
            blocks.append({"type": "table", "rows": rows})
            continue

        # Bullet list
        if re.match(r"^[-*+]\s+", stripped):
            items = []
            while i < len(lines) and re.match(r"^[-*+]\s+", lines[i].strip()):
                items.append(re.sub(r"^[-*+]\s+", "", lines[i].strip()))
                i += 1
            blocks.append({"type": "bullet_list", "items": items})
            continue

        # Ordered list
        if re.match(r"^\d+\.\s+", stripped):
            items = []
            while i < len(lines) and re.match(r"^\d+\.\s+", lines[i].strip()):
                items.append(re.sub(r"^\d+\.\s+", "", lines[i].strip()))
                i += 1
            blocks.append({"type": "ordered_list", "items": items})
            continue

        # Paragraph (collect until blank or special)
        para_lines = [stripped]
        i += 1
        while i < len(lines):
            nxt = lines[i].strip()
            if not nxt or nxt.startswith("#") or "|" in nxt or re.match(r"^[-*+\d]", nxt):
                break
            para_lines.append(nxt)
            i += 1
        blocks.append({"type": "paragraph", "text": " ".join(para_lines)})

    return blocks


def _apply_inline_formatting(paragraph: Any, text: str) -> None:
    """Apply bold/italic from ** and * markers."""
    pattern = re.compile(r"(\*\*(.+?)\*\*|\*(.+?)\*|([^*]+))")
    for match in pattern.finditer(text):
        if match.group(2):
            run = paragraph.add_run(match.group(2))
            run.bold = True
        elif match.group(3):
            run = paragraph.add_run(match.group(3))
            run.italic = True
        elif match.group(4):
            paragraph.add_run(match.group(4))


def build_docx(
    markdown_body: str,
    template_path: Path,
    output_path: Path,
) -> Path:
    """Convert markdown body to docx using template styles."""
    template_path = Path(template_path)
    output_path = Path(output_path)

    if template_path.exists():
        doc = Document(str(template_path))
    else:
        doc = Document()

    styles = _get_or_create_styles(doc)
    blocks = _parse_markdown_blocks(markdown_body)

    for block in blocks:
        btype = block["type"]

        if btype in ("h1", "h2", "h3", "paragraph"):
            style_name = styles.get(btype, "Normal")
            p = doc.add_paragraph(style=style_name)
            _apply_inline_formatting(p, block["text"])

        elif btype == "bullet_list":
            for item in block["items"]:
                p = doc.add_paragraph(style=styles["bullet_list"])
                _apply_inline_formatting(p, item)

        elif btype == "ordered_list":
            for item in block["items"]:
                p = doc.add_paragraph(style=styles["ordered_list"])
                _apply_inline_formatting(p, item)

        elif btype == "table":
            rows = block["rows"]
            if not rows:
                continue
            cols = max(len(r) for r in rows)
            table: Table = doc.add_table(rows=len(rows), cols=cols)
            table.style = styles["table"]
            for ri, row in enumerate(rows):
                for ci in range(cols):
                    cell_text = row[ci] if ci < len(row) else ""
                    table.rows[ri].cells[ci].text = cell_text

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


def create_default_template(output_path: Path) -> Path:
    """Create a minimal default.docx template with standard styles."""
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    doc.add_heading("Document Title", level=0)
    doc.add_heading("Heading 1 Example", level=1)
    doc.add_heading("Heading 2 Example", level=2)
    doc.add_heading("Heading 3 Example", level=3)
    doc.add_paragraph("Normal body text example.")
    doc.add_paragraph("Bullet item example.", style="List Bullet")
    doc.add_paragraph("Numbered item example.", style="List Number")

    table = doc.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Header A"
    table.rows[0].cells[1].text = "Header B"
    table.rows[1].cells[0].text = "Cell A"
    table.rows[1].cells[1].text = "Cell B"

    # Clear body content but keep styles — save then reload pattern:
    # python-docx keeps styles in the saved file.
    doc.save(str(output_path))
    return output_path
